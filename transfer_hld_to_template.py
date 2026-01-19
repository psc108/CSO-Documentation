#!/usr/bin/env python3
"""
Transfer HLD content into SecureCloud template
Preserves template structure and formatting
"""

import sys
import os
from pathlib import Path

# Check if running in venv, if not create instructions
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed")
    print("\nTo run this script:")
    print("1. python3 -m venv /tmp/docx_env")
    print("2. source /tmp/docx_env/bin/activate")
    print("3. pip install python-docx")
    print("4. python3 transfer_hld_to_template.py")
    sys.exit(1)

def find_paragraph_by_text(doc, search_text):
    """Find paragraph containing specific text"""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            return i
    return None

def insert_content_after(doc, after_index, content_lines, style_name='Normal'):
    """Insert content after a specific paragraph index"""
    # Insert in reverse to maintain order
    for line in reversed(content_lines):
        para = doc.paragraphs[after_index]._element
        new_para = para.getparent().insert(para.getparent().index(para) + 1, 
                                           doc.add_paragraph(line, style=style_name)._element)

def replace_placeholder(doc, old_text, new_text):
    """Replace placeholder text in document"""
    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)

def main():
    # Paths
    template_path = "/home/scottp/IdeaProjects/CSO-Documentation/docx/Templates/SecureCloud_service_HLD_xxx_Template_v0.1.docx"
    hld_md_path = "/home/scottp/IdeaProjects/CSO-Documentation/HLD.md"
    output_path = "/home/scottp/IdeaProjects/CSO-Documentation/docx/HLD_SecureCloud_Formatted.docx"
    
    print("Loading template...")
    doc = Document(template_path)
    
    print("Reading HLD content...")
    with open(hld_md_path, 'r') as f:
        hld_content = f.read()
    
    # Update placeholders
    print("Updating template placeholders...")
    replace_placeholder(doc, "[Subject]", "UK MSVX CPS Infrastructure")
    replace_placeholder(doc, "[Category]", "Infrastructure")
    replace_placeholder(doc, "[Status]", "Draft")
    replace_placeholder(doc, "[Publish Date]", "2024-12-20")
    replace_placeholder(doc, "[Comments]", "1.0")
    
    # Find section markers
    print("Locating template sections...")
    intro_idx = find_paragraph_by_text(doc, "Introduction")
    
    if intro_idx:
        print(f"Found Introduction at paragraph {intro_idx}")
        
        # Add content after Introduction
        intro_content = [
            "This document describes the high-level architecture for the UK MSVX CPS (Chief Security Officer Shared Services Portal) infrastructure deployment on AWS.",
            "",
            "The solution provides a scalable, highly available multi-tier application platform supporting:",
            "• Identity and Access Management via OpenStack Keystone",
            "• Service Catalog and Orchestration capabilities",
            "• Message Queuing and Event Processing through RabbitMQ",
            "• Administrative and User Interfaces for portal management",
            "• Reporting and Analytics for operational insights"
        ]
        
        # Insert content
        for i, line in enumerate(intro_content):
            para = doc.add_paragraph(line)
            # Move to correct position
            intro_para = doc.paragraphs[intro_idx]
            intro_para._element.addnext(para._element)
    
    print(f"Saving formatted document to: {output_path}")
    doc.save(output_path)
    print("✓ Document created successfully!")
    print("\nNOTE: This is a basic transfer. Manual review required for:")
    print("  - Complete content mapping")
    print("  - Table of contents update")
    print("  - Diagram insertion")
    print("  - Metadata completion")
    print("  - Style verification")

if __name__ == "__main__":
    main()
